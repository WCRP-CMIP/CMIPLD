#!/usr/bin/env python3
"""
Word Document Template Generator

Creates Word documents outlining form options, descriptions, and filling instructions
based on CSV field definitions and Python configuration files.

pip install python-docx

"""

import csv
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

failed_templates = []

def load_template_data(py_file):
    """Load configuration and data from Python file."""
    namespace = {}
    with open(py_file, 'r', encoding='utf-8') as f:
        exec(f.read(), namespace)
    
    config = namespace.get('TEMPLATE_CONFIG', {})
    data = namespace.get('DATA', {})
    return config, data

def load_csv_fields(csv_file):
    """Load field definitions from CSV."""
    fields = []
    with open(csv_file, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            fields.append(row)
    
    fields.sort(key=lambda x: int(x['field_order']))
    return fields

def setup_document_styles(doc):
    """Set up custom styles for the document."""
    
    # Title style
    try:
        title_style = doc.styles['Title']
    except KeyError:
        title_style = doc.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    # Heading 1 style
    try:
        h1_style = doc.styles['Heading 1']
    except KeyError:
        h1_style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = None  # Default color
    
    # Heading 2 style
    try:
        h2_style = doc.styles['Heading 2']
    except KeyError:
        h2_style = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = 'Calibri'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    
    # Field Label style
    try:
        field_label_style = doc.styles['Field Label']
    except KeyError:
        field_label_style = doc.styles.add_style('Field Label', WD_STYLE_TYPE.PARAGRAPH)
    field_label_style.font.name = 'Calibri'
    field_label_style.font.size = Pt(12)
    field_label_style.font.bold = True
    field_label_style.paragraph_format.space_after = Pt(6)
    
    # Field Description style
    try:
        field_desc_style = doc.styles['Field Description']
    except KeyError:
        field_desc_style = doc.styles.add_style('Field Description', WD_STYLE_TYPE.PARAGRAPH)
    field_desc_style.font.name = 'Calibri'
    field_desc_style.font.size = Pt(11)
    field_desc_style.paragraph_format.space_after = Pt(12)
    field_desc_style.paragraph_format.left_indent = Inches(0.25)

def add_field_to_document(doc, field_def, data, field_number):
    """Add a single field to the Word document."""
    
    field_type = field_def['field_type']
    field_id = field_def['field_id']
    label = field_def['label']
    description = field_def['description']
    data_source = field_def['data_source']
    required = field_def['required'].lower() == 'true'
    placeholder = field_def['placeholder']
    options_type = field_def['options_type']
    default_value = field_def['default_value']
    
    # Skip markdown fields as they're just informational
    if field_type == 'markdown':
        # Add markdown content as informational text
        if description:
            p = doc.add_paragraph()
            p.style = 'Field Description'
            formatted_desc = description.replace('\\n', '\n')
            p.add_run(formatted_desc)
            doc.add_paragraph()  # Add spacing
        return
    
    # Field header
    field_header = f"{field_number}. {label}"
    if required:
        field_header += " *"
    
    p = doc.add_paragraph(field_header)
    p.style = 'Field Label'
    
    # Field type and ID info
    type_info = f"Field Type: {field_type.title()}"
    if field_id:
        type_info += f" | Field ID: {field_id}"
    
    p = doc.add_paragraph(type_info)
    p.style = 'Field Description'
    run = p.runs[0]
    run.italic = True
    
    # Description
    if description and description.strip():
        formatted_desc = description.replace('\\n', '\n')
        p = doc.add_paragraph(f"Description: {formatted_desc}")
        p.style = 'Field Description'
    
    # Placeholder information
    if field_type in ['input', 'textarea'] and placeholder:
        p = doc.add_paragraph(f"Placeholder text: \"{placeholder}\"")
        p.style = 'Field Description'
    
    # Default value
    if default_value:
        p = doc.add_paragraph(f"Default value: {default_value}")
        p.style = 'Field Description'
    
    # Options for dropdowns and multi-selects
    if field_type in ['dropdown', 'multi-select']:
        options = get_field_options(field_def, data)
        
        if options:
            if field_type == 'multi-select':
                p = doc.add_paragraph("Available options (multiple selections allowed):")
            else:
                p = doc.add_paragraph("Available options:")
            p.style = 'Field Description'
            
            # Create a bulleted list of options
            for option in options:
                p = doc.add_paragraph(f"• {option}")
                p.style = 'Field Description'
                p.paragraph_format.left_indent = Inches(0.5)
    
    # Filling instructions
    instructions = get_filling_instructions(field_type, required, field_def, data)
    if instructions:
        p = doc.add_paragraph(f"How to fill: {instructions}")
        p.style = 'Field Description'
        run = p.runs[0]
        run.font.bold = True
    
    # Add spacing between fields
    doc.add_paragraph()

def get_field_options(field_def, data):
    """Extract options for dropdown and multi-select fields."""
    data_source = field_def['data_source']
    options_type = field_def['options_type']
    field_id = field_def['field_id']
    
    options = []
    
    if data_source != 'none' and data_source in data:
        source_data = data[data_source]
        
        if options_type == 'dict_keys':
            options = list(source_data.keys())
        elif options_type == 'list':
            options = list(source_data)
        elif options_type in ['dict_multiple']:
            options = list(source_data.keys())
        elif options_type == 'dict_with_extra':
            options = list(source_data.keys())
            options.extend(["Open Source", "Registration Required", "Proprietary"])
        elif options_type == 'list_with_na':
            options = ["Not applicable"] + list(source_data)
    
    return options

def get_filling_instructions(field_type, required, field_def, data):
    """Generate filling instructions based on field type."""
    
    instructions = []
    
    if field_type == 'input':
        instructions.append("Enter a single line of text")
    elif field_type == 'textarea':
        instructions.append("Enter multiple lines of text")
    elif field_type == 'dropdown':
        instructions.append("Select one option from the dropdown list")
    elif field_type == 'multi-select':
        instructions.append("Select multiple options from the list (hold Ctrl/Cmd to select multiple)")
    elif field_type == 'checkboxes':
        instructions.append("Check one or more boxes")
    
    if required:
        instructions.append("This field is required and must be filled")
    else:
        instructions.append("This field is optional")
    
    # Add specific instructions based on field content
    placeholder = field_def.get('placeholder', '')
    if placeholder:
        instructions.append(f"Use the placeholder as guidance: '{placeholder}'")
    
    return ". ".join(instructions) + "."

def generate_word_document(config, fields, data):
    """Generate complete Word document."""
    
    doc = Document()
    setup_document_styles(doc)
    
    # Document title
    title = doc.add_paragraph(config['name'])
    title.style = 'Title'
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Document description
    if config.get('description'):
        desc_p = doc.add_paragraph(config['description'])
        desc_p.style = 'Field Description'
        desc_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
    doc.add_paragraph()  # Add spacing
    
    # Overview section
    overview = doc.add_paragraph("Overview")
    overview.style = 'Heading 1'
    
    overview_text = f"""This document outlines the form fields and options for the "{config['name']}" template. 
Each field is numbered and includes information about its type, available options, and instructions for completion.

Template Information:
• Name: {config['name']}
• Labels: {config.get('labels', 'N/A')}
• Total Fields: {len([f for f in fields if f['field_type'] != 'markdown'])}
• Required Fields: {len([f for f in fields if f['required'].lower() == 'true' and f['field_type'] != 'markdown'])}
"""
    
    p = doc.add_paragraph(overview_text)
    p.style = 'Field Description'
    
    doc.add_paragraph()
    
    # Fields section
    fields_header = doc.add_paragraph("Form Fields")
    fields_header.style = 'Heading 1'
    
    field_number = 1
    for field_def in fields:
        # Merge config defaults into data if needed
        if field_def['field_id'] not in data and field_def['field_id'] in config:
            data[field_def['field_id']] = config[field_def['field_id']]
        
        add_field_to_document(doc, field_def, data, field_number)
        
        # Only increment counter for non-markdown fields
        if field_def['field_type'] != 'markdown':
            field_number += 1
    
    # Footer section
    doc.add_page_break()
    footer_header = doc.add_paragraph("Additional Information")
    footer_header.style = 'Heading 1'
    
    footer_text = """Completion Guidelines:
• Fields marked with an asterisk (*) are required
• For dropdown fields, select the most appropriate option
• For multi-select fields, you may choose multiple options
• For text fields, provide clear and concise information
• If unsure about a field, refer to the description provided

For technical support or questions about this template, please refer to the project documentation or contact the development team."""
    
    p = doc.add_paragraph(footer_text)
    p.style = 'Field Description'
    
    return doc

def process_template_pair(template_name, csv_file, py_file, output_dir):
    """Process a single template pair and generate Word document."""
    
    print(f"Processing {template_name}...")
    
    try:
        config, data = load_template_data(py_file)
        if not config:
            print(f"    No TEMPLATE_CONFIG found")
            return False
        
        print(f"    Loaded: {config['name']}")
        
        fields = load_csv_fields(csv_file)
        print(f"    Fields: {len(fields)}")
        
        doc = generate_word_document(config, fields, data)
        
        # Save the document
        output_file = output_dir / f"{template_name}_guide.docx"
        doc.save(str(output_file))
        
        print(f"    Generated {template_name}_guide.docx")
        return True
        
    except Exception as e:
        print(f"    Error: {e}")
        failed_templates.append(f"{template_name}: {str(e)}")
        return False

def parse_arguments():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(description='Generate Word document guides for issue templates')
    parser.add_argument('-t', '--template-dir', type=Path, help='Template directory')
    parser.add_argument('-o', '--output-dir', type=Path, help='Output directory')
    parser.add_argument('--template', type=str, help='Generate guide for specific template')
    return parser.parse_args()

def main():
    """Main function."""
    args = parse_arguments()
    
    print("Word Document Template Generator")
    print("Generating user guides for issue templates")
    
    script_dir = Path.cwd()
    template_dir = args.template_dir or script_dir / ".github" / "GEN_ISSUE_TEMPLATE"
    output_dir = args.output_dir or script_dir / "docs" / "template_guides_docx"
    
    print(f"Template dir: {template_dir}")
    print(f"Output dir: {output_dir}")
    
    if not template_dir.exists():
        print(f"Template directory not found: {template_dir}")
        return False
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    csv_files = list(template_dir.glob('*.csv'))
    
    if args.template:
        csv_files = [f for f in csv_files if f.stem == args.template]
    
    if not csv_files:
        print("No CSV files found")
        return False
    
    print(f"Processing {len(csv_files)} templates")
    
    success_count = 0
    for csv_file in csv_files:
        template_name = csv_file.stem
        py_file = template_dir / f"{template_name}.py"
        
        if py_file.exists():
            if process_template_pair(template_name, csv_file, py_file, output_dir):
                success_count += 1
        else:
            print(f"    Warning: No Python config file found for {template_name}")
    
    print(f"Results: {success_count}/{len(csv_files)} successful")
    
    if failed_templates:
        print("\nFailed templates:")
        for failure in failed_templates:
            print(f"  - {failure}")
    
    return success_count > 0

if __name__ == '__main__':
    try:
        success = main()
        sys.exit(0 if success else 1)
    except ImportError as e:
        if 'docx' in str(e):
            print("Error: python-docx package not found. Please install it:")
            print("pip install python-docx")
            sys.exit(1)
        else:
            raise
